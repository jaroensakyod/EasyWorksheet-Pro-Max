# docx_exporter.py - Word document generation using python-docx
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

class DocxExporter:
    """Export worksheets to Word document format"""
    
    def create_worksheet_doc(self, title, school_name, topic, questions, answers):
        """Create a Word document worksheet"""
        doc = Document()
        
        # Title
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # School name
        if school_name:
            school_para = doc.add_paragraph(school_name)
            school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Topic
        topic_para = doc.add_paragraph(f"หัวข้อ: {topic}")
        topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Questions section
        doc.add_heading("แบบฝึกหัด / Exercises", level=1)
        
        for i, question in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {question}")
        
        doc.add_paragraph()
        
        # Answer key section
        doc.add_heading("เฉลย / Answer Key", level=1)
        
        for i, answer in enumerate(answers, 1):
            doc.add_paragraph(f"{i}. {answer}")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def create_wordsearch_doc(self, title, school_name, topic, grid, placed_words, answers=None):
        """Create a Word document with word search"""
        doc = Document()
        
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if school_name:
            school_para = doc.add_paragraph(school_name)
            school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        topic_para = doc.add_paragraph(f"หัวข้อ: {topic}")
        topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Word search grid
        doc.add_heading("ปริศนาหาคำ / Word Search", level=1)
        
        # Add grid as preformatted text
        grid_text = "\n".join([" ".join(row) for row in grid])
        grid_para = doc.add_paragraph(grid_text)
        grid_para.style = "Normal"
        grid_para.paragraph_format.font.name = "Courier New"
        grid_para.paragraph_format.font.size = Pt(12)
        
        # Words to find
        doc.add_heading("คำศัพท์ที่ต้องหา / Words to Find:", level=1)
        for word in placed_words:
            doc.add_paragraph(word)
        
        doc.add_paragraph()
        
        # Answer key
        if answers:
            doc.add_heading("เฉลย / Answer Key", level=1)
            for word in answers:
                doc.add_paragraph(word)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def create_tracing_doc(self, title, school_name, topic, lines):
        """Create a Word document for tracing practice"""
        doc = Document()
        
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if school_name:
            school_para = doc.add_paragraph(school_name)
            school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"หัวข้อ: {topic}")
        doc.add_paragraph()
        
        doc.add_heading("แบบฝึกคัดลายมือ / Tracing Practice", level=1)
        
        for text in lines:
            # Large text for tracing
            p = doc.add_paragraph()
            runner = p.add_run(text)
            runner.font.size = Pt(36)
            runner.font.name = "TH Sarabun New"
            
            doc.add_paragraph()  # Spacing
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

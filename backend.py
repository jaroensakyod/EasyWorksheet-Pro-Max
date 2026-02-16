# backend.py (Updated with PDF/Docx Extraction and Quiz Gen)
# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import random
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib.colors import black, lightgrey, blue
import io
import qrcode
from reportlab.lib.utils import ImageReader
import google.generativeai as genai
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from pypdf import PdfReader

# Import Groq for Groq support
try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

# Import OpenRouter for OpenRouter support
try:
    import openai
    OPENROUTER_AVAILABLE = True
except ImportError:
    OPENROUTER_AVAILABLE = False

class WorksheetGenerator:
    def __init__(self, ai_api_key=None, provider="Google Gemini"):
        self.model = None
        self.model_name = None
        self.provider = provider
        self.client = None  # For Groq/OpenRouter
        
        if ai_api_key and provider == "Google Gemini":
            try:
                genai.configure(api_key=ai_api_key)
                
                # Smart Model Detection (Free Tier Compatible Models)
                model_priority = [
                    'gemini-1.5-flash',      # Best for Free Tier - fast & quota-friendly
                    'gemini-1.5-pro',       # Free Tier available
                    'gemini-2.0-flash-exp', # New model (if available)
                ]
                
                try:
                    # List available models
                    available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                    print(f"Available models: {available_models}")
                    
                    # Find first matching model from priority list
                    for model_name in model_priority:
                        for available in available_models:
                            if model_name in available:
                                self.model = genai.GenerativeModel(available)
                                self.model_name = available
                                print(f"[OK] Using AI Model: {available}")
                                break
                        if self.model:
                            break
                    
                    # If still no model, use first available (fallback)
                    if not self.model and available_models:
                        # Prefer flash model if available
                        for m in available_models:
                            if 'flash' in m.lower():
                                self.model = genai.GenerativeModel(m)
                                self.model_name = m
                                print(f"[OK] Using Flash Model (fallback): {m}")
                                break
                        
                        # If no flash found, use first available
                        if not self.model:
                            self.model = genai.GenerativeModel(available_models[0])
                            self.model_name = available_models[0]
                            print(f"[OK] Using first available model: {available_models[0]}")
                            
                except Exception as e:
                    print(f"[!] Model detection failed: {e}")
                    # Last resort fallback - try flash explicitly
                    try:
                        self.model = genai.GenerativeModel('gemini-1.5-flash')
                        self.model_name = 'gemini-1.5-flash'
                        print(f"[OK] Using Flash as fallback")
                    except Exception as fallback_error:
                        print(f"[!] Fallback also failed: {fallback_error}")
                        self.model = None
                        
            except Exception as e:
                print(f"[!] Google API configuration failed: {e}")
                self.model = None
                
        elif ai_api_key and provider == "Groq":
            if GROQ_AVAILABLE:
                try:
                    self.client = Groq(api_key=ai_api_key)
                    # Groq uses Llama and other models - fast inference
                    self.model_name = "llama-3.3-70b-versatile"  # Will be set per request
                    print(f"[OK] Groq client initialized successfully")
                except Exception as e:
                    print(f"[!] Groq client initialization failed: {e}")
                    self.client = None
            else:
                print("[!] Groq library not installed. Install with: pip install groq")
                
        elif ai_api_key and provider == "OpenRouter":
            if OPENROUTER_AVAILABLE:
                try:
                    self.client = openai.OpenAI(
                        api_key=ai_api_key,
                        base_url="https://openrouter.ai/api/v1"
                    )
                    self.model_name = "openrouter-llama"  # Will be set per request
                    print(f"[OK] OpenRouter client initialized successfully")
                except Exception as e:
                    print(f"[!] OpenRouter client initialization failed: {e}")
                    self.client = None
            else:
                print("[!] OpenRouter (openai compatible) not properly configured.")

        # Register Thai Font for PDF (Fallback logic)
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        self.font_name = 'Helvetica'
        try:
            pdfmetrics.registerFont(TTFont('ThaiFont', 'C:/Windows/Fonts/tahoma.ttf'))
            self.font_name = 'ThaiFont'
        except:
            try:
                pdfmetrics.registerFont(TTFont('ThaiFont', 'C:/Windows/Fonts/leelawad.ttf'))
                self.font_name = 'ThaiFont'
            except:
                pass

    def _generate_content(self, prompt):
        """Generate content using the appropriate provider."""
        if self.provider == "Google Gemini" and self.model:
            try:
                response = self.model.generate_content(prompt)
                return response.text
            except Exception as e:
                raise Exception(f"Google Gemini API error: {e}")
                
        elif self.provider == "Groq" and self.client:
            try:
                # Groq uses Llama models - fast and efficient
                chat_completion = self.client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "You are a helpful Thai education assistant that creates worksheets and exercises."},
                        {"role": "user", "content": prompt}
                    ],
                    model="llama-3.3-70b-versatile",  # Latest Llama 3 model
                    temperature=0.7,
                )
                return chat_completion.choices[0].message.content
            except Exception as e:
                raise Exception(f"Groq API error: {e}")
                
        elif self.provider == "OpenRouter" and self.client:
            try:
                chat_completion = self.client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "You are a helpful Thai education assistant that creates worksheets and exercises."},
                        {"role": "user", "content": prompt}
                    ],
                    model="meta-llama/llama-3-8b-instruct",  # Llama 3 via OpenRouter
                    temperature=0.7,
                )
                return chat_completion.choices[0].message.content
            except Exception as e:
                raise Exception(f"OpenRouter API error: {e}")
        else:
            raise Exception(f"No valid API configuration for {self.provider}")

    def extract_text_from_file(self, uploaded_file):
        """Extracts text from PDF or Docx."""
        text = ""
        try:
            if uploaded_file.name.endswith(".pdf"):
                reader = PdfReader(uploaded_file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            elif uploaded_file.name.endswith(".docx"):
                doc = Document(uploaded_file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
        except Exception as e:
            return f"Error extracting text: {e}"
        return text

    def generate_ai_worksheet(self, topic, grade_level, num_questions=10):
        """Generates worksheets for non-calculation topics using AI."""
        if not self.model and not self.client:
            return ["Error: No API Key configured"], ["Please add API Key"]
        
        # Create detailed prompts for each topic
        prompts = {
            # ป.1
            "การวัดความยาว": f"Create {num_questions} length measurement exercises for Grade 1 students. Include: measuring lines, comparing lengths, drawing lines of specific lengths. Use simple units.",
            "การชั่ง": f"Create {num_questions} weighing exercises for Grade 1 students. Include: reading scales, comparing weights, simple word problems about weight.",
            "การตวง": f"Create {num_questions} volume measurement exercises for Grade 1 students. Include: measuring liquids, comparing volumes, using standard units.",
            "รูปเรขาคณิต": f"Create {num_questions} geometry exercises for Grade 1 students. Include: identifying shapes (circle, square, triangle, rectangle), drawing shapes, counting sides and corners.",
            "เวลา": f"Create {num_questions} time exercises for Grade 1 students. Include: reading analog clocks (o'clock), drawing clock hands, sequencing daily activities.",
            
            # ป.2
            "การวัดความยาว": f"Create {num_questions} length measurement exercises for Grade 2 students. Include: measuring in cm and m, comparing lengths, solving word problems.",
            "การชั่ง": f"Create {num_questions} weighing exercises for Grade 2 students. Include: reading scales in kg, comparing weights, simple word problems.",
            "เวลา": f"Create {num_questions} time exercises for Grade 2 students. Include: reading analog clocks (half hour, quarter hour), writing time, time word problems.",
            "เงิน": f"Create {num_questions} money exercises for Grade 2 students. Include: counting Thai baht, making change, money word problems.",
            "การตวง": f"Create {num_questions} volume measurement exercises for Grade 2 students. Include: measuring in liters, comparing volumes, word problems.",
            "รูปเรขาคณิต": f"Create {num_questions} geometry exercises for Grade 2 students. Include: identifying 2D shapes, counting sides and corners, drawing shapes.",
            
            # ป.3
            "แผนภูมิรูปภาพและแผนภูมิแท่ง": f"Create {num_questions} pictograph and bar chart exercises for Grade 3 students. Include: reading data, interpreting charts, creating charts from data.",
            "การวัดความยาว": f"Create {num_questions} length measurement exercises for Grade 3 students. Include: measuring in mm, cm, m, km, converting units, word problems.",
            "เวลา": f"Create {num_questions} time exercises for Grade 3 students. Include: reading analog clocks (5-minute intervals), 24-hour clock, elapsed time problems.",
            "การชั่ง การตวง": f"Create {num_questions} measurement exercises (weight and volume) for Grade 3 students. Include: kg, g, liters, ml, conversions, word problems.",
            "เงินและการบันทึกรายรับรายจ่าย": f"Create {num_questions} money and record-keeping exercises for Grade 3 students. Include: calculating totals, making change, income-expense records.",
            "จุด เส้นตรง รังสี ส่วนของเส้นตรง มุม": f"Create {num_questions} geometry exercises about points, lines, rays, line segments and angles for Grade 3. Include: identifying, naming, drawing.",
            "รูปเรขาคณิต": f"Create {num_questions} geometry exercises for Grade 3 students. Include: identifying 2D and 3D shapes, counting edges, faces, vertices.",
            
            # ป.4
            "เรขาคณิต": f"Create {num_questions} geometry exercises for Grade 4 students. Include: identifying 2D shapes, properties of shapes, perimeter, area basics.",
            "แผนภูมิรูปภาพ แผนภูมิแท่ง และตาราง": f"Create {num_questions} data representation exercises (pictographs, bar charts, tables) for Grade 4 students.",
            "การวัด": f"Create {num_questions} measurement exercises for Grade 4 students. Include: length, weight, volume, time units, conversions.",
            "พื้นที่": f"Create {num_questions} area exercises for Grade 4 students. Include: finding area by counting squares, area of rectangle, word problems.",
            "เงิน": f"Create {num_questions} money exercises for Grade 4 students. Include: complex transactions, discounts, savings, word problems.",
            "เวลา": f"Create {num_questions} time exercises for Grade 4 students. Include: 24-hour clock, time intervals, calendar problems.",
            
            # ป.5
            "มุม": f"Create {num_questions} angle exercises for Grade 5 students. Include: measuring angles with protractor, drawing angles, types of angles (acute, right, obtuse, straight).",
            "เส้นขนาน": f"Create {num_questions} parallel lines exercises for Grade 5 students. Include: identifying parallel lines, properties of parallel lines, drawing parallel lines.",
            "สถิติและความน่าจะเป็นเบื้องต้น": f"Create {num_questions} statistics and probability exercises for Grade 5 students. Include: mean, mode, median, simple probability.",
            "บทประยุกต์": f"Create {num_questions} mixed application problems for Grade 5 students combining various math topics.",
            "รูปสี่เหลี่ยม": f"Create {num_questions} quadrilateral exercises for Grade 5 students. Include: identifying types (square, rectangle, parallelogram, trapezoid), properties, perimeter, area.",
            "รูปสามเหลี่ยม": f"Create {num_questions} triangle exercises for Grade 5 students. Include: types of triangles, properties, drawing, area of triangle basics.",
            "รูปวงกลม": f"Create {num_questions} circle exercises for Grade 5 students. Include: identifying parts of circle, radius, diameter, circumference basics.",
            "รูปเรขาคณิตสามมิติและปริมาตรของทรงสี่เหลี่ยมมุมฉาก": f"Create {num_questions} 3D geometry and volume exercises for Grade 5. Include: cube, rectangular prism, volume calculation.",
            
            # ป.6
            "ตัวประกอบของจำนวนนับ": f"Create {num_questions} factors and multiples exercises for Grade 6 students. Include: finding factors, prime factors, greatest common factor, least common multiple.",
            "เส้นขนาน": f"Create {num_questions} parallel lines exercises for Grade 6 students. Include: properties of parallel lines, angle relationships, transversal.",
            "สมการและการแก้สมการ": f"Create {num_questions} equation solving exercises for Grade 6 students. Include: one-step equations, two-step equations, word problems with variables.",
            "ทิศ แผนที่และแผนผัง": f"Create {num_questions} directions, maps and plans exercises for Grade 6 students. Include: 8 compass directions, reading maps, scale drawing basics.",
            "รูปสี่เหลี่ยม": f"Create {num_questions} quadrilateral exercises for Grade 6. Include: all types, properties, area of various quadrilaterals.",
            "รูปวงกลม": f"Create {num_questions} circle exercises for Grade 6. Include: radius, diameter, circumference, area of circle.",
            "บทประยุกต์": f"Create {num_questions} complex application problems for Grade 6 combining various math topics.",
            "รูปเรขาคณิตสามมิติและปริมาตรของทรงสี่เหลี่ยมมุมฉาก": f"Create {num_questions} 3D geometry and volume exercises for Grade 6. Include: surface area, volume of various prisms.",
            "สถิติและความน่าจะเป็นเบื้องต้น": f"Create {num_questions} statistics and probability exercises for Grade 6. Include: data collection, graphs, probability experiments.",
        }
        
        # Get prompt for topic or use default
        prompt = prompts.get(topic, f"Create {num_questions} math exercises about '{topic}' for {grade_level} students.")
        
        # Add grade context
        grade_context = {
            "ป.1": "Grade 1 (Thailand IPST Curriculum)",
            "ป.2": "Grade 2 (Thailand IPST Curriculum)",
            "ป.3": "Grade 3 (Thailand IPST Curriculum)",
            "ป.4": "Grade 4 (Thailand IPST Curriculum)",
            "ป.5": "Grade 5 (Thailand IPST Curriculum)",
            "ป.6": "Grade 6 (Thailand IPST Curriculum)",
        }
        
        full_prompt = f"""
        {prompt}
        
        Target Grade: {grade_context.get(grade_level, grade_level)}
        Curriculum: IPST (Thailand Institute of Scientific and Technological Research)
        
        Output format:
        Q: [Question]
        A: [Answer]
        
        Create exactly {num_questions} questions.
        """
        
        try:
            resp_text = self._generate_content(full_prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    def generate_quiz_from_text(self, text, num_questions=5, grade="General"):
        """Generates a multiple choice quiz from text using AI."""
        if not self.model and not self.client:
            return ["Error: No API Key configured"], ["Please add API Key"]
        
        # Limit text to avoid token limits (approx first 2000 chars)
        context_text = text[:3000] 
        
        prompt = f"""
        Create a {num_questions}-question multiple choice quiz (in Thai) based on this text:
        "{context_text}"
        
        Target audience: {grade}
        Format exactly like this for each question:
        Q: [Question]
        A: [Option A]
        B: [Option B]
        C: [Option C]
        D: [Option D]
        Ans: [Correct Letter]
        """
        try:
            content = self._generate_content(prompt)
            # Parse is tricky, let's keep it simple string list for PDF now
            # We will split by "Q:"
            questions = []
            answers = [] # We can store correct answers separately if needed
            
            raw_qs = content.split("Q:")
            for raw in raw_qs:
                if not raw.strip(): continue
                lines = raw.strip().split('\n')
                q_text = lines[0].strip()
                options = []
                correct = ""
                for line in lines[1:]:
                    if line.startswith("Ans:"): correct = line.replace("Ans:", "").strip()
                    elif line.strip(): options.append(line.strip())
                
                if q_text:
                    full_q = q_text + "\n" + "\n".join(options)
                    questions.append(full_q)
                    answers.append(correct)
            
            return questions, answers
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    def generate_questions(self, operation, num_questions, min_val, max_val):
        # ... (Existing Math Logic) ...
        questions = []
        answers = []
        for _ in range(num_questions):
            a = random.randint(min_val, max_val)
            b = random.randint(min_val, max_val)
            question_text = ""
            answer_val = 0
            if operation == "Addition (+)":
                question_text = f"{a} + {b} = ____"
                answer_val = a + b
            elif operation == "Subtraction (-)":
                if a < b: a, b = b, a
                question_text = f"{a} - {b} = ____"
                answer_val = a - b
            elif operation == "Multiplication (x)":
                question_text = f"{a} x {b} = ____"
                answer_val = a * b
            elif operation == "Division (÷)":
                answer_val = random.randint(min_val, max_val)
                a = answer_val * b
                question_text = f"{a} ÷ {b} = ____"
            questions.append(question_text)
            answers.append(str(answer_val))
        return questions, answers

    def generate_ai_word_problems(self, topic, grade_level, num_questions):
        """Generates math word problems using AI."""
        if not self.model and not self.client:
            return ["No API Key configured"], ["Please add API Key"]
        prompt = f"Generate {num_questions} math word problems (Thai) for {grade_level} about '{topic}'. Output format:\nQ: [Question]\nA: [Answer]"
        try:
            resp_text = self._generate_content(prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"Error: {str(e)}"], ["Error"]

    def generate_science_worksheet(self, topic, grade_level, num_questions=10):
        """Generates science worksheets using AI based on IPST curriculum."""
        if not self.model and not self.client:
            return ["Error: No API Key configured"], ["Please add API Key"]
        
        # Detailed prompts for science topics
        science_prompts = {
            # ป.1
            "สิ่งต่างรอบตัวเรา (สิ่งมีชีวิต, สิ่งไม่มีชีวิต, สมบัติของวัสดุ)": f"Create {num_questions} science exercises about living things, non-living things, and properties of materials for Grade 1. Include: identifying living/non-living, basic properties of materials.",
            "พืชรอบตัวเรา (ส่วนประกอบของพืช, การเจริญเติบโต)": f"Create {num_questions} science exercises about plants for Grade 1. Include: parts of plants, plant growth, plant needs.",
            "สัตว์รอบตัวเรา (สัตว์หลากหลายชนิด, การดูแลสัตว์)": f"Create {num_questions} science exercises about animals for Grade 1. Include: types of animals, animal characteristics, caring for pets.",
            "ดวงดาวและท้องฟ้า (ดวงอาทิตย์, ดวงจันทร์, ดวงดาว)": f"Create {num_questions} science exercises about the sun, moon, and stars for Grade 1. Include: basic astronomy concepts, observations of the sky.",
            "สภาพอากาศ (หนาว, ร้อน, ฝน, ลม)": f"Create {num_questions} science exercises about weather (cold, hot, rain, wind) for Grade 1. Include: weather observation, types of weather.",
            
            # ป.2
            "สิ่งมีชีวิตกับการดำรงชีวิต (อาหาร, ที่อยู่อาศัย, การสืบพันธุ์)": f"Create {num_questions} science exercises about how living things survive for Grade 2. Include: food, shelter, reproduction of plants and animals.",
            "สิ่งแวดล้อม (แสง, เสียง, ความร้อน)": f"Create {num_questions} science exercises about the environment for Grade 2. Include: light, sound, heat - basic properties and sources.",
            "น้ำและอากาศ (สถานะของน้ำ, การเกิดฝน)": f"Create {num_questions} science exercises about water and air for Grade 2. Include: states of water, water cycle, rain formation.",
            "ดิน (องค์ประกอบของดิน, ชนิดของดิน)": f"Create {num_questions} science exercises about soil for Grade 2. Include: components of soil, types of soil, soil uses.",
            "ท้องฟ้าและการพยากรณ์อากาศ (การสังเกตเมฆ, การพยากรณ์อากาศ)": f"Create {num_questions} science exercises about sky observation and weather prediction for Grade 2. Include: cloud observation, simple weather forecasting.",
            
            # ป.3
            "ร่างกายของเรา (ระบบย่อยอาหาร, ระบบหายใจ)": f"Create {num_questions} science exercises about the human body for Grade 3. Include: digestive system, respiratory system, basic anatomy.",
            "พืชกับการดำรงชีวิต (การสังเคราะห์ด้วยแสง, การขยายพันธุ์)": f"Create {num_questions} science exercises about plants for Grade 3. Include: photosynthesis, plant reproduction, plant life cycle.",
            "สิ่งมีชีวิตกับสิ่งแวดล้อม (ห่วงโซ่อาหาร, สมดุลในธรรมชาติ)": f"Create {num_questions} science exercises about food chains and natural balance for Grade 3. Include: ecosystems, food chains, ecological balance.",
            "วัสดุรอบตัว (โลหะ, ไม้, พลาสติก)": f"Create {num_questions} science exercises about materials around us for Grade 3. Include: properties of metal, wood, plastic, uses of materials.",
            "แรงและการเคลื่อนที่ (แรงผลัก, แรงดึง, แรงเสียดทาน)": f"Create {num_questions} science exercises about forces and motion for Grade 3. Include: push/pull forces, friction, basic motion concepts.",
            "พลังงาน (ความร้อน, แสง, เสียง)": f"Create {num_questions} science exercises about energy for Grade 3. Include: heat, light, sound - sources and properties.",
            
            # ป.4
            "ระบบร่างกาย (ระบบหมุนเวียนเลือด, ระบบขับถ่าย)": f"Create {num_questions} science exercises about human body systems for Grade 4. Include: circulatory system, excretory system.",
            "พืชที่หลากหลาย (การจำแนกพืช, การสืบพันธุ์พืช)": f"Create {num_questions} science exercises about plant diversity for Grade 4. Include: plant classification, plant reproduction methods.",
            "สิ่งมีชีวิตกับสิ่งแวดล้อม (แหล่งน้ำ, ห่วงโซ่อาหาร)": f"Create {num_questions} science exercises about living things and environment for Grade 4. Include: water sources, food chains in ecosystems.",
            "สสาร (สถานะ, การเปลี่ยนแปลง)": f"Create {num_questions} science exercises about matter for Grade 4. Include: states of matter, physical changes.",
            "แรงและความดัน (แรงในธรรมชาติ, ความดันอากาศ)": f"Create {num_questions} science exercises about forces and pressure for Grade 4. Include: natural forces, air pressure.",
            "พลังงานไฟฟ้า (ไฟฟ้าพื้นฐาน, วงจรไฟฟ้า)": f"Create {num_questions} science exercises about electrical energy for Grade 4. Include: basics of electricity, electric circuits.",
            
            # ป.5
            "ระบบสุขภาพ (ฮอร์โมน, การเจริญเติบโต)": f"Create {num_questions} science exercises about health system for Grade 5. Include: hormones, human growth and development.",
            "การสืบพันธุ์ (การสืบพันธุ์สัตว์, การสืบพันธุ์พืช)": f"Create {num_questions} science exercises about reproduction for Grade 5. Include: animal reproduction, plant reproduction.",
            "สิ่งแวดล้อม (การถ่ายทอดพลังงาน, สิ่งมีชีวิตกับสิ่งแวดล้อม)": f"Create {num_questions} science exercises about environment for Grade 5. Include: energy transfer, living things and their environment.",
            "สสาร (อะตอม, ธาตุ, สารประกอบ)": f"Create {num_questions} science exercises about matter for Grade 5. Include: atoms, elements, compounds.",
            "แรงและการเคลื่อนที่ (แรงโน้มถ่วง, แรงเสียดทาน)": f"Create {num_questions} science exercises about forces and motion for Grade 5. Include: gravity, friction, motion.",
            "คลื่น (คลื่นเสียง, คลื่นแสง)": f"Create {num_questions} science exercises about waves for Grade 5. Include: sound waves, light waves, wave properties.",
            
            # ป.6
            "ระบบต่อมไร้ท่อ (ฮอร์โมน, ต่อมไร้ท่อสำคัญ)": f"Create {num_questions} science exercises about endocrine system for Grade 6. Include: hormones, major endocrine glands.",
            "พันธุศาสตร์เบื้องต้น (ลักษณะทางพันธุกรรม, การถ่ายทอดลักษณะ)": f"Create {num_questions} science exercises about genetics for Grade 6. Include: inherited traits, genetic traits inheritance.",
            "วิวัฒนาการ (การเปลี่ยนแปลงของสิ่งมีชีวิต)": f"Create {num_questions} science exercises about evolution for Grade 6. Include: changes in living things over time.",
            "สสารและพลังงาน (กฎทรงพลังงาน, การถ่ายโอนพลังงาน)": f"Create {num_questions} science exercises about matter and energy for Grade 6. Include: law of conservation of energy, energy transfer.",
            "ระบบสุริยะ (ดาวเคราะห์, การเกิดกลางวัน-กลางคืน)": f"Create {num_questions} science exercises about the solar system for Grade 6. Include: planets, day-night cycle, solar system.",
            "สิ่งแวดล้อม (ทรัพยากรธรรมชาติ, การอนุรักษ์)": f"Create {num_questions} science exercises about environment for Grade 6. Include: natural resources, conservation.",
            
            # ม.1
            "สารบริสุทธิ์": f"Create {num_questions} science exercises about pure substances for Grade 7 (ม.1). Include: elements, compounds, mixtures, separation techniques.",
            "หน่วยพื้นฐานของสิ่งมีชีวิต": f"Create {num_questions} science exercises about basic units of living things for Grade 7 (ม.1). Include: cells, cell structure, cell function.",
            "หน่วยพื้นฐานของการดำรงชีวิตของพืช": f"Create {num_questions} science exercises about plant biology basics for Grade 7 (ม.1). Include: plant cells, photosynthesis, plant structures.",
            "พลังงานความร้อน": f"Create {num_questions} science exercises about heat energy for Grade 7 (ม.1). Include: heat transfer, thermal energy, temperature.",
            "กระบวนการเปลี่ยนแปลงลมฟ้าอากาศ": f"Create {num_questions} science exercises about weather and climate changes for Grade 7 (ม.1). Include: weather patterns, climate, atmospheric changes.",
            
            # ม.2
            "สารละลาย": f"Create {num_questions} science exercises about solutions for Grade 8 (ม.2). Include: solubility, concentration, types of solutions.",
            "ร่างกายมนุษย์": f"Create {num_questions} science exercises about the human body for Grade 8 (ม.2). Include: body systems, organs, human anatomy.",
            "การเคลื่อนที่และแรง": f"Create {num_questions} science exercises about motion and forces for Grade 8 (ม.2). Include: Newton's laws, velocity, acceleration.",
            "งานและพลังงาน": f"Create {num_questions} science exercises about work and energy for Grade 8 (ม.2). Include: work, kinetic/potential energy, energy transformation.",
            "การแยกสาร": f"Create {num_questions} science exercises about separation of substances for Grade 8 (ม.2). Include: filtration, distillation, chromatography.",
            "โลกและการเปลี่ยนแปลง": f"Create {num_questions} science exercises about Earth and its changes for Grade 8 (ม.2). Include: geological processes, plate tectonics, Earth's history.",
            
            # ม.3
            "พันธุศาสตร์": f"Create {num_questions} science exercises about genetics for Grade 9 (ม.3). Include: Mendelian genetics, DNA, heredity patterns.",
            "คลื่นและแสง": f"Create {num_questions} science exercises about waves and light for Grade 9 (ม.3). Include: wave properties, light reflection/refraction, optics.",
            "ระบบสุริยะของเรา": f"Create {num_questions} science exercises about our solar system for Grade 9 (ม.3). Include: planets, asteroids, comets, space exploration.",
            "ปฏิกิริยาเคมีและวัสดุในชีวิตประจำวัน": f"Create {num_questions} science exercises about chemical reactions and everyday materials for Grade 9 (ม.3). Include: types of reactions, common chemicals, materials science.",
            "ไฟฟ้า": f"Create {num_questions} science exercises about electricity for Grade 9 (ม.3). Include: circuits, Ohm's law, electrical energy, electromagnetism.",
            "ระบบนิเวศและความหลากหลายทางชีวภาพ": f"Create {num_questions} science exercises about ecosystems and biodiversity for Grade 9 (ม.3). Include: ecological relationships, biodiversity, conservation.",
            
            # ===== เคมี (Chemistry) ม.4-6 =====
            # ม.4
            "อะตอมและสมบัติของธาตุ": f"Create {num_questions} chemistry exercises about atoms and properties of elements for Grade 10 (ม.4). Include: atomic structure, periodic table, electron configuration, periodic trends.",
            "พันธะเคมี": f"Create {num_questions} chemistry exercises about chemical bonding for Grade 10 (ม.4). Include: ionic bonding, covalent bonding, metallic bonding, Lewis structures.",
            "ปริมาณสัมพันธ์ในปฏิกิริยาเคมี": f"Create {num_questions} chemistry exercises about stoichiometry for Grade 10 (ม.4). Include: balancing equations, mole concept, limiting reactants, yield calculations.",
            # ม.5
            "สมบัติของก๊าซและสมการเคมี": f"Create {num_questions} chemistry exercises about gas properties and chemical equations for Grade 11 (ม.5). Include: gas laws, ideal gas equation, types of chemical equations.",
            "อัตราการเกิดปฏิกิริยาเคมี": f"Create {num_questions} chemistry exercises about reaction rates for Grade 11 (ม.5). Include: rate laws, activation energy, factors affecting reaction rates, reaction mechanisms.",
            "สมดุลเคมี": f"Create {num_questions} chemistry exercises about chemical equilibrium for Grade 11 (ม.5). Include: equilibrium constant, Le Chatelier's principle, equilibrium calculations.",
            "กรด-เบส": f"Create {num_questions} chemistry exercises about acids and bases for Grade 11 (ม.5). Include: pH scale, acid-base reactions, buffers, titration.",
            # ม.6
            "ไฟฟ้าเคมี": f"Create {num_questions} chemistry exercises about electrochemistry for Grade 12 (ม.6). Include: galvanic cells, electrolysis, standard reduction potentials, Faraday's laws.",
            "ธาตุอินทรีย์และสารชีวโมเลกุล": f"Create {num_questions} chemistry exercises about organic compounds and biomolecules for Grade 12 (ม.6). Include: hydrocarbons, functional groups, carbohydrates, proteins, lipids, nucleic acids.",
            "เคมีอินทรีย์": f"Create {num_questions} chemistry exercises about organic chemistry for Grade 12 (ม.6). Include: organic reactions, synthesis, spectroscopy, polymer chemistry.",
            
            # ===== ฟิสิกส์ (Physics) ม.4-6 =====
            # ม.4
            "การเคลื่อนที่แนวตรง": f"Create {num_questions} physics exercises about linear motion for Grade 10 (ม.4). Include: displacement, velocity, acceleration, kinematic equations, free fall.",
            "แรงและกฎการเคลื่อนที่": f"Create {num_questions} physics exercises about forces and laws of motion for Grade 10 (ม.4). Include: Newton's three laws, friction, equilibrium, inclined planes.",
            "งานและพลังงาน": f"Create {num_questions} physics exercises about work and energy for Grade 10 (ม.4). Include: work, kinetic energy, potential energy, conservation of energy, power.",
            "โมเมนตัมและการชน": f"Create {num_questions} physics exercises about momentum and collisions for Grade 10 (ม.4). Include: momentum conservation, elastic/inelastic collisions, center of mass.",
            # ม.5
            "การเคลื่อนที่ในระบบต่างๆ": f"Create {num_questions} physics exercises about motion in various systems for Grade 11 (ม.5). Include: circular motion, projectile motion, simple harmonic motion.",
            "แรงในธรรมชาติ": f"Create {num_questions} physics exercises about natural forces for Grade 11 (ม.5). Include: gravitational force, electrostatic force, magnetic force, universal gravitation.",
            "คลื่น": f"Create {num_questions} physics exercises about waves for Grade 11 (ม.5). Include: wave properties, wave equation, superposition, standing waves.",
            "เสียง": f"Create {num_questions} physics exercises about sound for Grade 11 (ม.5). Include: sound waves, speed of sound, Doppler effect, resonance, intensity.",
            "แสง": f"Create {num_questions} physics exercises about light for Grade 11 (ม.5). Include: reflection, refraction, lenses, mirrors, optical instruments, interference, diffraction.",
            # ม.6
            "ไฟฟ้าสถิตและไฟฟ้ากระแส": f"Create {num_questions} physics exercises about static and current electricity for Grade 12 (ม.6). Include: electric fields, Coulomb's law, circuits, Ohm's law, Kirchhoff's laws.",
            "แม่เหล็กไฟฟ้า": f"Create {num_questions} physics exercises about electromagnetism for Grade 12 (ม.6). Include: magnetic fields, electromagnetic induction, Faraday's law, Maxwell's equations, AC circuits.",
            "ฟิสิกส์อะตอม": f"Create {num_questions} physics exercises about atomic physics for Grade 12 (ม.6). Include: atomic structure, quantum theory, photoelectric effect, Bohr model, atomic spectra.",
            "ฟิสิกส์นิวเคลียร์": f"Create {num_questions} physics exercises about nuclear physics for Grade 12 (ม.6). Include: nuclear structure, radioactivity, nuclear reactions, fission, fusion, half-life.",
            
            # ===== ชีววิทยา (Biology) ม.4-6 =====
            # ม.4
            "ระบบย่อยอาหาร": f"Create {num_questions} biology exercises about digestive system for Grade 10 (ม.4). Include: digestive organs, enzymes, absorption, nutrition.",
            "ระบบหายใจ": f"Create {num_questions} biology exercises about respiratory system for Grade 10 (ม.4). Include: respiratory organs, gas exchange, cellular respiration, breathing mechanism.",
            "ระบบหมุนเวียนเลือด": f"Create {num_questions} biology exercises about circulatory system for Grade 10 (ม.4). Include: heart, blood vessels, blood composition, blood circulation, lymphatic system.",
            "ระบบขับถ่าย": f"Create {num_questions} biology exercises about excretory system for Grade 10 (ม.4). Include: kidneys, urine formation, waste elimination, osmoregulation.",
            "ระบบประสาท": f"Create {num_questions} biology exercises about nervous system for Grade 10 (ม.4). Include: neurons, nerve impulses, brain, spinal cord, reflexes, senses.",
            "ระบบต่อมไร้ท่อ": f"Create {num_questions} biology exercises about endocrine system for Grade 10 (ม.4). Include: hormones, endocrine glands, feedback mechanisms, homeostasis.",
            # ม.5
            "การถ่ายทอดสารภายในร่างกาย": f"Create {num_questions} biology exercises about transport in living organisms for Grade 11 (ม.5). Include: transport in plants, transport in animals, circulation, diffusion, osmosis, active transport.",
            "ระบบภูมิคุ้มกัน": f"Create {num_questions} biology exercises about immune system for Grade 11 (ม.5). Include: innate immunity, adaptive immunity, antibodies, vaccines, allergic reactions.",
            "การสืบพันธุ์และพัฒนาการ": f"Create {num_questions} biology exercises about reproduction and development for Grade 11 (ม.5). Include: reproductive systems, fertilization, embryonic development, birth process, growth.",
            "การถ่ายทอดลักษณะทางพันธุกรรม": f"Create {num_questions} biology exercises about inheritance for Grade 11 (ม.5). Include: Mendelian genetics, Mendel's laws, genetic crosses, probability, sex-linked inheritance.",
            # ม.6
            "พันธุศาสตร์": f"Create {num_questions} biology exercises about genetics for Grade 12 (ม.6). Include: DNA structure, replication, gene expression, mutations, genetic engineering, biotechnology.",
            "พันธุกรรมเทคโนโลยี": f"Create {num_questions} biology exercises about genetic technology for Grade 12 (ม.6). Include: recombinant DNA, PCR, gel electrophoresis, cloning, GMOs, gene therapy.",
            "วิวัฒนาการ": f"Create {num_questions} biology exercises about evolution for Grade 12 (ม.6). Include: natural selection, adaptation, speciation, evidence for evolution, human evolution.",
            "นิเวศวิทยา": f"Create {num_questions} biology exercises about ecology for Grade 12 (ม.6). Include: ecosystems, energy flow, nutrient cycles, ecological succession, population dynamics.",
            "สิ่งแวดล้อม": f"Create {num_questions} biology exercises about environment for Grade 12 (ม.6). Include: biodiversity, conservation, pollution, climate change, sustainable development.",
        }
        
        # Get prompt for topic or use default
        prompt = science_prompts.get(topic, f"Create {num_questions} science exercises about '{topic}' for {grade_level} students according to Thailand IPST curriculum.")
        
        full_prompt = f"""
        {prompt}
        
        Target Grade/Level: {grade_level}
        Curriculum: IPST (Thailand Institute of Scientific and Technological Research)
        
        Output format:
        Q: [Question in Thai]
        A: [Answer in Thai]
        
        Create exactly {num_questions} questions.
        Include a mix of: multiple choice, fill in the blank, short answer, and true/false questions.
        """
        
        try:
            resp_text = self._generate_content(full_prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    def generate_chemistry_worksheet(self, topic, grade_level, num_questions=10):
        """Generates chemistry worksheets using AI based on IPST curriculum for ม.4-6."""
        if not self.model and not self.client:
            return ["Error: No API Key configured"], ["Please add API Key"]
        
        chemistry_prompts = {
            # ม.4
            "อะตอมและสมบัติของธาตุ": f"Create {num_questions} chemistry exercises about atoms and properties of elements for Grade 10 (ม.4). Include: atomic structure, periodic table, electron configuration, periodic trends, atomic number, mass number.",
            "พันธะเคมี": f"Create {num_questions} chemistry exercises about chemical bonding for Grade 10 (ม.4). Include: ionic bonding, covalent bonding, metallic bonding, Lewis structures, VSEPR theory, electronegativity.",
            "ปริมาณสัมพันธ์ในปฏิกิริยาเคมี": f"Create {num_questions} chemistry exercises about stoichiometry for Grade 10 (ม.4). Include: balancing chemical equations, mole concept, molar mass, limiting reactants, theoretical yield, percent yield.",
            # ม.5
            "สมบัติของก๊าซและสมการเคมี": f"Create {num_questions} chemistry exercises about gas properties and chemical equations for Grade 11 (ม.5). Include: Boyle's law, Charles's law, Avogadro's law, ideal gas equation, PV=nRT, types of chemical equations.",
            "อัตราการเกิดปฏิกิริยาเคมี": f"Create {num_questions} chemistry exercises about reaction rates for Grade 11 (ม.5). Include: rate laws, rate constants, activation energy, Arrhenius equation, factors affecting reaction rates, reaction mechanisms, catalysts.",
            "สมดุลเคมี": f"Create {num_questions} chemistry exercises about chemical equilibrium for Grade 11 (ม.5). Include: equilibrium constant (Kc, Kp), Le Chatelier's principle, equilibrium calculations, homogeneous/heterogeneous equilibrium.",
            "กรด-เบส": f"Create {num_questions} chemistry exercises about acids and bases for Grade 11 (ม.5). Include: pH scale, pOH, acid-base theories, strong/weak acids and bases, neutralization reactions, titration curves, buffers, hydrolysis.",
            # ม.6
            "ไฟฟ้าเคมี": f"Create {num_questions} chemistry exercises about electrochemistry for Grade 12 (ม.6). Include: galvanic cells, electrolytic cells, standard reduction potentials, cell potential, Faraday's laws of electrolysis, corrosion, batteries.",
            "ธาตุอินทรีย์และสารชีวโมเลกุล": f"Create {num_questions} chemistry exercises about organic compounds and biomolecules for Grade 12 (ม.6). Include: hydrocarbons, functional groups, carbohydrates (monosaccharides, disaccharides, polysaccharides), proteins, lipids, nucleic acids (DNA, RNA).",
            "เคมีอินทรีย์": f"Create {num_questions} chemistry exercises about organic chemistry for Grade 12 (ม.6). Include: organic reaction mechanisms, substitution, addition, elimination reactions, polymerization, organic synthesis, IR and NMR spectroscopy basics.",
        }
        
        grade_context = {
            "ม.4": "Grade 10 / Matthayom 4 (Thailand IPST Chemistry Curriculum)",
            "ม.5": "Grade 11 / Matthayom 5 (Thailand IPST Chemistry Curriculum)",
            "ม.6": "Grade 12 / Matthayom 6 (Thailand IPST Chemistry Curriculum)",
        }
        
        prompt = chemistry_prompts.get(topic, f"Create {num_questions} chemistry exercises about '{topic}' for {grade_level} students according to Thailand IPST curriculum.")
        
        full_prompt = f"""
        {prompt}
        
        Target Grade/Level: {grade_context.get(grade_level, grade_level)}
        Curriculum: IPST (Thailand Institute of Scientific and Technological Research) Chemistry
        
        Output format:
        Q: [Question in Thai]
        A: [Answer in Thai]
        
        Create exactly {num_questions} questions.
        Include a mix of: multiple choice, fill in the blank, short answer, calculation problems, and true/false questions.
        For calculation problems, show the solution steps clearly.
        """
        
        try:
            resp_text = self._generate_content(full_prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    def generate_physics_worksheet(self, topic, grade_level, num_questions=10):
        """Generates physics worksheets using AI based on IPST curriculum for ม.4-6."""
        if not self.model and not self.client:
            return ["Error: No API Key configured"], ["Please add API Key"]
        
        physics_prompts = {
            # ม.4
            "การเคลื่อนที่แนวตรง": f"Create {num_questions} physics exercises about linear motion for Grade 10 (ม.4). Include: displacement, velocity, acceleration, kinematic equations, free fall, motion graphs (position-time, velocity-time).",
            "แรงและกฎการเคลื่อนที่": f"Create {num_questions} physics exercises about forces and laws of motion for Grade 10 (ม.4). Include: Newton's three laws, types of forces (gravity, normal, tension, friction), free body diagrams, equilibrium, inclined planes.",
            "งานและพลังงาน": f"Create {num_questions} physics exercises about work and energy for Grade 10 (ม.4). Include: work done by forces, kinetic energy, gravitational potential energy, conservation of mechanical energy, power, efficiency.",
            "โมเมนตัมและการชน": f"Create {num_questions} physics exercises about momentum and collisions for Grade 10 (ม.4). Include: linear momentum, conservation of momentum, elastic collisions, inelastic collisions, center of mass, impulse.",
            # ม.5
            "การเคลื่อนที่ในระบบต่างๆ": f"Create {num_questions} physics exercises about motion in various systems for Grade 11 (ม.5). Include: uniform circular motion, centripetal force, projectile motion, simple harmonic motion, damped and driven oscillations.",
            "แรงในธรรมชาติ": f"Create {num_questions} physics exercises about natural forces for Grade 11 (ม.5). Include: gravitational force, Coulomb's law, magnetic force, universal gravitation, orbital motion, satellite motion.",
            "คลื่น": f"Create {num_questions} physics exercises about waves for Grade 11 (ม.5). Include: wave characteristics (amplitude, wavelength, frequency, period), wave types (longitudinal, transverse), wave equation v=fλ, superposition principle, standing waves.",
            "เสียง": f"Create {num_questions} physics exercises about sound for Grade 11 (ม.5). Include: sound wave properties, speed of sound, Doppler effect, intensity and loudness (decibels), resonance, beats, acoustic phenomena.",
            "แสง": f"Create {num_questions} physics exercises about light for Grade 11 (ม.5). Include: reflection (plane/spherical mirrors), refraction (Snell's law), lenses (converging/diverging), optical instruments (microscope, telescope), interference, diffraction, polarization basics.",
            # ม.6
            "ไฟฟ้าสถิตและไฟฟ้ากระแส": f"Create {num_questions} physics exercises about static and current electricity for Grade 12 (ม.6). Include: electric field, electric potential, capacitance, Ohm's law, resistivity, series/parallel circuits, Kirchhoff's laws, electric power.",
            "แม่เหล็กไฟฟ้า": f"Create {num_questions} physics exercises about electromagnetism for Grade 12 (ม.6). Include: magnetic field around conductors, Ampere's law, electromagnetic induction, Faraday's law, Lenz's law, transformers, AC circuits, RL/RC/LC circuits.",
            "ฟิสิกส์อะตอม": f"Create {num_questions} physics exercises about atomic physics for Grade 12 (ม.6). Include: atomic models (Rutherford, Bohr), quantum theory, photoelectric effect, Compton scattering, atomic spectra, X-rays, wave-particle duality.",
            "ฟิสิกส์นิวเคลียร์": f"Create {num_questions} physics exercises about nuclear physics for Grade 12 (ม.6). Include: nuclear structure, nuclear forces, radioactivity (alpha, beta, gamma decay), radioactive decay laws (half-life), nuclear reactions (fission, fusion), nuclear energy, radiation safety.",
        }
        
        grade_context = {
            "ม.4": "Grade 10 / Matthayom 4 (Thailand IPST Physics Curriculum)",
            "ม.5": "Grade 11 / Matthayom 5 (Thailand IPST Physics Curriculum)",
            "ม.6": "Grade 12 / Matthayom 6 (Thailand IPST Physics Curriculum)",
        }
        
        prompt = physics_prompts.get(topic, f"Create {num_questions} physics exercises about '{topic}' for {grade_level} students according to Thailand IPST curriculum.")
        
        full_prompt = f"""
        {prompt}
        
        Target Grade/Level: {grade_context.get(grade_level, grade_level)}
        Curriculum: IPST (Thailand Institute of Scientific and Technological Research) Physics
        
        Output format:
        Q: [Question in Thai]
        A: [Answer in Thai]
        
        Create exactly {num_questions} questions.
        Include a mix of: multiple choice, fill in the blank, short answer, calculation problems, and diagram-based questions.
        For calculation problems, show the solution steps clearly with formulas used.
        """
        
        try:
            resp_text = self._generate_content(full_prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    def generate_biology_worksheet(self, topic, grade_level, num_questions=10):
        """Generates biology worksheets using AI based on IPST curriculum for ม.4-6."""
        if not self.model and not self.client:
            return ["Error: No API Key configured"], ["Please add API Key"]
        
        biology_prompts = {
            # ม.4 - ระบบต่างๆ ในร่างกาย
            "ระบบย่อยอาหาร": f"Create {num_questions} biology exercises about digestive system for Grade 10 (ม.4). Include: digestive organs (mouth, esophagus, stomach, small/large intestine), digestive enzymes, mechanical/chemical digestion, absorption, nutrients, digestive disorders.",
            "ระบบหายใจ": f"Create {num_questions} biology exercises about respiratory system for Grade 10 (ม.4). Include: respiratory organs, breathing mechanism, gas exchange in alveoli, cellular respiration (aerobic/anaerobic), respiratory diseases, oxygen transport.",
            "ระบบหมุนเวียนเลือด": f"Create {num_questions} biology exercises about circulatory system for Grade 10 (ม.4). Include: heart structure and function, blood vessels (arteries, veins, capillaries), blood composition, blood circulation (systemic/pulmonary), heart cycle, blood pressure, cardiovascular diseases.",
            "ระบบขับถ่าย": f"Create {num_questions} biology exercises about excretory system for Grade 10 (ม.4). Include: kidney structure, nephron function, urine formation (filtration, reabsorption, secretion), osmoregulation, excretory organs, waste elimination, kidney diseases.",
            "ระบบประสาท": f"Create {num_questions} biology exercises about nervous system for Grade 10 (ม.4). Include: neuron structure, nerve impulse transmission, synapse, brain (cerebrum, cerebellum, brainstem), spinal cord, reflexes, senses (vision, hearing, touch, taste, smell), nervous system disorders.",
            "ระบบต่อมไร้ท่อ": f"Create {num_questions} biology exercises about endocrine system for Grade 10 (ม.4). Include: endocrine glands (pituitary, thyroid, parathyroid, adrenal, pancreas, gonads), hormone functions, feedback mechanisms, homeostasis, endocrine disorders, hormone interactions.",
            # ม.5
            "การถ่ายทอดสารภายในร่างกาย": f"Create {num_questions} biology exercises about transport in living organisms for Grade 11 (ม.5). Include: xylem and phloem transport in plants, transpiration, translocation, blood circulation in animals, lymphatic system, diffusion, osmosis, active transport across cell membranes.",
            "ระบบภูมิคุ้มกัน": f"Create {num_questions} biology exercises about immune system for Grade 11 (ม.5). Include: innate immunity (physical, chemical, cellular barriers), adaptive immunity (B cells, T cells), antibodies (immunoglobulins), vaccines, immune responses, allergies, autoimmune diseases, immunodeficiency.",
            "การสืบพันธุ์และพัฒนาการ": f"Create {num_questions} biology exercises about reproduction and development for Grade 11 (ม.5). Include: male/female reproductive systems, gametogenesis, fertilization, embryonic development (cleavage, gastrulation, organogenesis), fetal development, birth process, hormonal regulation, reproductive technologies.",
            "การถ่ายทอดลักษณะทางพันธุกรรม": f"Create {num_questions} biology exercises about inheritance for Grade 11 (ม.5). Include: Mendel's laws of inheritance, monohybrid and dihybrid crosses, Punnett squares, probability in genetics, incomplete dominance, codominance, multiple alleles, sex-linked inheritance, genetic counseling basics.",
            # ม.6
            "พันธุศาสตร์": f"Create {num_questions} biology exercises about genetics for Grade 12 (ม.6). Include: DNA structure and replication, RNA transcription, genetic code, translation (protein synthesis), gene expression regulation, mutations (point mutations, chromosomal mutations), genetic disorders, gene therapy.",
            "พันธุกรรมเทคโนโลยี": f"Create {num_questions} biology exercises about genetic technology for Grade 12 (ม.6). Include: recombinant DNA technology, restriction enzymes, DNA ligase, PCR (polymerase chain reaction), gel electrophoresis, DNA sequencing, cloning (reproductive/therapeutic), GMOs, CRISPR-Cas9, ethical issues in genetic engineering.",
            "วิวัฒนาการ": f"Create {num_questions} biology exercises about evolution for Grade 12 (ม.6). Include: Darwin's theory of natural selection, evidence for evolution (fossils, comparative anatomy, embryology, molecular biology), speciation, adaptive radiation, convergent/divergent evolution, coevolution, human evolution, extinction.",
            "นิเวศวิทยา": f"Create {num_questions} biology exercises about ecology for Grade 12 (ม.6). Include: ecosystems (biotic/abiotic factors), energy flow (food chains, food webs, trophic levels), biogeochemical cycles (carbon, nitrogen, water), ecological succession, population ecology (growth models, carrying capacity), community interactions, biodiversity indices.",
            "สิ่งแวดล้อม": f"Create {num_questions} biology exercises about environment for Grade 12 (ม.6). Include: biodiversity importance, threats to biodiversity (habitat loss, pollution, overexploitation, invasive species), conservation strategies (protected areas, wildlife corridors, captive breeding), pollution (air, water, soil), climate change impacts, sustainable development, environmental policies.",
        }
        
        grade_context = {
            "ม.4": "Grade 10 / Matthayom 4 (Thailand IPST Biology Curriculum)",
            "ม.5": "Grade 11 / Matthayom 5 (Thailand IPST Biology Curriculum)",
            "ม.6": "Grade 12 / Matthayom 6 (Thailand IPST Biology Curriculum)",
        }
        
        prompt = biology_prompts.get(topic, f"Create {num_questions} biology exercises about '{topic}' for {grade_level} students according to Thailand IPST curriculum.")
        
        full_prompt = f"""
        {prompt}
        
        Target Grade/Level: {grade_context.get(grade_level, grade_level)}
        Curriculum: IPST (Thailand Institute of Scientific and Technological Research) Biology
        
        Output format:
        Q: [Question in Thai]
        A: [Answer in Thai]
        
        Create exactly {num_questions} questions.
        Include a mix of: multiple choice, fill in the blank, short answer, diagram identification, and true/false questions.
        Use scientific terminology appropriate for Thai high school biology students.
        """
        
        try:
            resp_text = self._generate_content(full_prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    def generate_word_search(self, words, grid_size=15):
        # ... (Existing Word Search Logic with grid_size) ...
        size = grid_size
        grid = [[' ' for _ in range(size)] for _ in range(size)]
        placed_words = []
        words = [w.strip().upper() for w in words if w.strip()]
        for word in words:
            word_clean = word.replace(" ", "")
            if len(word_clean) > size: continue
            placed = False; attempts = 0
            while not placed and attempts < 100:
                d = random.choice([(0,1),(1,0),(1,1)])
                r = random.randint(0,size-1); c = random.randint(0,size-1)
                er = r + d[0]*(len(word_clean)-1); ec = c + d[1]*(len(word_clean)-1)
                if 0<=er<size and 0<=ec<size:
                    can = True
                    for i in range(len(word_clean)):
                        if grid[r+d[0]*i][c+d[1]*i] not in [' ', word_clean[i]]: can = False; break
                    if can:
                        for i in range(len(word_clean)): grid[r+d[0]*i][c+d[1]*i] = word_clean[i]
                        placed_words.append(word); placed = True
                attempts += 1
        letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        for r in range(size):
            for c in range(size):
                if grid[r][c] == ' ': grid[r][c] = random.choice(letters)
        return grid, placed_words

    def generate_tracing_lines(self, text_input):
        lines = []
        words = text_input.split(",")
        for word in words:
            w = word.strip()
            lines.append((w + "   ") * (5 if len(w)<10 else 2))
        return lines

    def generate_thai_worksheet(self, topic, grade_level, num_questions=10, exercise_type="mix"):
        """Generates Thai language worksheets using AI based on Thai curriculum."""
        if not self.model:
            return ["Error: No AI Key"], ["Please add API Key"]
        
        # Grade context mapping
        grade_context = {
            "ป.1": "ประถมศึกษาปีที่ 1 (Grade 1)",
            "ป.2": "ประถมศึกษาปีที่ 2 (Grade 2)",
            "ป.3": "ประถมศึกษาปีที่ 3 (Grade 3)",
            "ป.4": "ประถมศึกษาปีที่ 4 (Grade 4)",
            "ป.5": "ประถมศึกษาปีที่ 5 (Grade 5)",
            "ป.6": "ประถมศึกษาปีที่ 6 (Grade 6)",
            "ม.1": "มัธยมศึกษาปีที่ 1 / ม.1 (Grade 7)",
            "ม.2": "มัธยมศึกษาปีที่ 2 / ม.2 (Grade 8)",
            "ม.3": "มัธยมศึกษาปีที่ 3 / ม.3 (Grade 9)",
            "ม.4": "มัธยมศึกษาปีที่ 4 / ม.4 (Grade 10)",
            "ม.5": "มัธยมศึกษาปีที่ 5 / ม.5 (Grade 11)",
            "ม.6": "มัธยมศึกษาปีที่ 6 / ม.6 (Grade 12)",
        }
        
        # Exercise type prompts
        exercise_prompts = {
            "mix": f"Create {num_questions} Thai language exercises covering various aspects including writing, reading, grammar, vocabulary, and literature appropriate for {grade_context.get(grade_level, grade_level)}.",
            "writing": f"Create {num_questions} Thai writing exercises (การเขียน) for {grade_context.get(grade_level, grade_level)}. Include: sentence writing, paragraph writing, creative writing, and composition exercises.",
            "reading": f"Create {num_questions} Thai reading comprehension exercises (การอ่าน) for {grade_context.get(grade_level, grade_level)}. Include: short passages, reading for understanding, and answering questions.",
            "grammar": f"Create {num_questions} Thai grammar exercises (หลักภาษา) for {grade_context.get(grade_level, grade_level)}. Include: parts of speech, sentence structure, verb conjugation, and language usage.",
            "vocabulary": f"Create {num_questions} Thai vocabulary exercises (คำศัพท์) for {grade_context.get(grade_level, grade_level)}. Include: word meanings, synonyms, antonyms, word usage, and vocabulary building.",
            "literature": f"Create {num_questions} Thai literature exercises (วรรณคดี) for {grade_context.get(grade_level, grade_level)}. Include: reading Thai poems, understanding Thai literature, and literary analysis.",
        }
        
        # Topic-specific prompts for better accuracy
        topic_prompts = {
            # ป.1
            "ตัวอักษรไทย (พยัญชนะไทย 44 ตัว, สระ 32 รูป)": f"Create {num_questions} Thai alphabet exercises for Grade 1. Cover: 44 Thai consonants, 32 vowel forms, reading and writing Thai letters.",
            "สระในภาษาไทย (สระเดี่ยว, สระประสม)": f"Create {num_questions} Thai vowel exercises for Grade 1. Cover: single vowels (สระเดี่ยว), combined vowels (สระประสม), vowel placement.",
            "การอ่านออกเสียง (อ่านคาบวรรณยุกต์)": f"Create {num_questions} Thai tone reading exercises for Grade 1. Cover: tone marks (วรรณยุกต์), reading with tones, tone practice.",
            "คำศัพท์พื้นฐาน (คำสิ่งของ, คำสัตว์, คำครอบครัว)": f"Create {num_questions} basic Thai vocabulary exercises for Grade 1. Cover: objects, animals, family members, everyday vocabulary.",
            "ประโยคและเรื่องสั้น (ประโยคสั้น, นิทานสั้น)": f"Create {num_questions} Thai sentence and short story exercises for Grade 1. Cover: simple sentences, short story reading, story comprehension.",
            # ป.2
            "คำและความหมาย (คำซ้ำ, คำตรงข้าม, คำพ้อง)": f"Create {num_questions} Thai word meaning exercises for Grade 2. Cover: reduplicated words (คำซ้ำ), antonyms (คำตรงข้าม), synonyms (คำพ้อง).",
            "หน่วยคำสรรพนาม (สรรพนาม, คำสรรพนามสรรพบุรณ)": f"Create {num_questions} Thai pronoun exercises for Grade 2. Cover: personal pronouns, demonstrative pronouns, reflexive pronouns.",
            "การเขียน (เขียนตามคำบอก, เขียนประโยค)": f"Create {num_questions} Thai writing exercises for Grade 2. Cover: dictation, sentence writing, guided writing.",
            "นิทานพื้นบ้าน (นิทานชาดก, นิทานพื้นบ้านไทย)": f"Create {num_questions} Thai folk tale exercises for Grade 2. Cover: Jataka stories, Thai folk tales, story reading and comprehension.",
            "การอ่านจับใจความ (อ่านเรื่องสั้น, ตอบคำถาม)": f"Create {num_questions} Thai reading comprehension exercises for Grade 2. Cover: short passages, finding main idea, answering questions.",
            # ป.3
            "ชนิดของคำ (คำนาม, คำกริยา, คำคุณศัพท์)": f"Create {num_questions} Thai parts of speech exercises for Grade 3. Cover: nouns, verbs, adjectives, identifying word types.",
            "กลอนแปด (โครงสร้างกลอนแปด, คำครุ-ลหุ)": f"Create {num_questions} Thai eight-syllable verse (กลอนแปด) exercises for Grade 3. Cover: structure of กลอนแปด, heavy-light syllables (คำครุ-ลหุ).",
            "การเขียนเรียงความ (เขียนเรียงความสั้น)": f"Create {num_questions} Thai composition exercises for Grade 3. Cover: short essay writing, paragraph development, expressing ideas.",
            "คำราชาศัพท์เบื้องต้น (คำขึ้นต้น-ลงท้าย)": f"Create {num_questions} Thai royal vocabulary (ราชาศัพท์) exercises for Grade 3. Cover: honorific words, opening and closing phrases.",
            "วรรณคดีไทย (ขุนช้างขุนแผน, สุภาษิตไทย)": f"Create {num_questions} Thai literature exercises for Grade 3. Cover: Khun Chang Khun Paen, Thai proverbs (สุภาษิต), literary appreciation.",
            # ป.4
            "หน่วยคำและความหมาย (คำภาษาต่างประเทศ, คำยืม)": f"Create {num_questions} Thai word formation exercises for Grade 4. Cover: loanwords, foreign words, word formation.",
            "ชนิดของคำ (คำสรรพนาม, คำสันธาน, คำบุพบท)": f"Create {num_questions} Thai parts of speech exercises for Grade 4. Cover: pronouns, conjunctions, prepositions.",
            "การอ่านตีความ (อ่านบทความ, อ่านข่าว)": f"Create {num_questions} Thai reading interpretation exercises for Grade 4. Cover: article reading, news reading, interpretation skills.",
            "การเขียนจดหมาย (จดหมายขอบคุณ, จดหมายเชิญ)": f"Create {num_questions} Thai letter writing exercises for Grade 4. Cover: thank you letters, invitation letters, formal letter writing.",
            "กลอนสุภาพ (โครงสร้างกลอนสุภาพ)": f"Create {num_questions} Thai Suphab verse (กลอนสุภาพ) exercises for Grade 4. Cover: structure of กลอนสุภาพ, verse composition.",
            # ป.5
            "ประโยคและองค์ประกอบ (องค์ประโยค, ชนิดของประโยค)": f"Create {num_questions} Thai sentence structure exercises for Grade 5. Cover: sentence components, types of sentences.",
            "วลีและอนุประโยค (วลีนาม, วลีกริยา)": f"Create {num_questions} Thai phrases and clauses exercises for Grade 5. Cover: noun phrases, verb phrases, clause structure.",
            "การเขียนรายงาน (รายงานการศึกษา, รายงานข่าว)": f"Create {num_questions} Thai report writing exercises for Grade 5. Cover: study reports, news reports, structured writing.",
            "วรรณคดีสุนทรียภาพ (กาพย์กลอนบทร้อยกรอง)": f"Create {num_questions} Thai aesthetic literature exercises for Grade 5. Cover: กาพย์, กลอน, poetic verses, literary beauty.",
            "ภาษาถิ่น (ภาษาอีสาน, ภาษาเหนือ, ภาษาใต้)": f"Create {num_questions} Thai regional dialect exercises for Grade 5. Cover: Isan dialect, Northern dialect, Southern dialect.",
            # ป.6
            "หลักการใช้คำ (คำราชาศัพท์, คำสุภาพ)": f"Create {num_questions} Thai language etiquette exercises for Grade 6. Cover: royal vocabulary, polite language, proper word usage.",
            "การเขียนเชิงสร้างสรรค์ (เรียงความ, นิทานสั้น)": f"Create {num_questions} Thai creative writing exercises for Grade 6. Cover: essays, short stories, creative expression.",
            "การอ่านวิเคราะห์ (วิเคราะห์เรื่อง, วิเคราะห์ข่าว)": f"Create {num_questions} Thai analytical reading exercises for Grade 6. Cover: story analysis, news analysis, critical reading.",
            "วรรณคดีวรรณกรรม (วรรณกรรมระดับชาติ)": f"Create {num_questions} Thai national literature exercises for Grade 6. Cover: national literature works, literary criticism, appreciation.",
            "การนำเสนอ (การพูด, การนำเสนอข้อมูล)": f"Create {num_questions} Thai presentation exercises for Grade 6. Cover: public speaking, data presentation, oral communication.",
            # ม.1
            "หน่วยคำสรรพนาม (การใช้สรรพนามในบริบทต่างๆ)": f"Create {num_questions} Thai pronoun usage exercises for ม.1. Cover: pronouns in different contexts, formal/informal usage.",
            "การเปลี่ยนรูปคำ (การผันคำกริยา, การลดรูปคำ)": f"Create {num_questions} Thai word formation exercises for ม.1. Cover: verb conjugation, word reduction, morphological changes.",
            "วลีและอนุประโยค (วลีขยาย, อนุประโยค)": f"Create {num_questions} Thai phrases and clauses exercises for ม.1. Cover: expanded phrases, subordinate clauses.",
            "วรรณคดี (ร้อยกรองไทย, กาพย์ยานเอก)": f"Create {num_questions} Thai poetry exercises for ม.1. Cover: Thai rhyming verses, กาพย์ยานเอก, poetic forms.",
            "การอ่าน-เขียน (อ่านบทความ, เขียนเรียงความ)": f"Create {num_questions} Thai reading and writing exercises for ม.1. Cover: article reading, essay writing, integrated skills.",
            # ม.2
            "คำและประโยคซ้อน (ประโยคซ้อน, ประโยคซ้อนกลบ)": f"Create {num_questions} Thai complex sentence exercises for ม.2. Cover: compound sentences, complex sentences, sentence embedding.",
            "กลอนแปด-กลอนสุภาพ (การแต่งกลอน, สัมผัสกลอน)": f"Create {num_questions} Thai verse composition exercises for ม.2. Cover: กลอนแปด, กลอนสุภาพ, rhyme schemes.",
            "วรรณคดีอีสาน (ลิเก, โขน, หนังใหญ่)": f"Create {num_questions} Thai regional literature exercises for ม.2. Cover: Likay, Khon, Nang Yai (Isan performing arts).",
            "การเขียนเชิงสร้างสรรค์ (เขียนนิยายสั้น, บทละคร)": f"Create {num_questions} Thai creative writing exercises for ม.2. Cover: short stories, play scripts, dramatic writing.",
            "ภาษาถิ่นและภาษากลาง (ความแตกต่าง, การใช้)": f"Create {num_questions} Thai dialect vs central Thai exercises for ม.2. Cover: differences between dialects, appropriate usage.",
            # ม.3
            "ภาษากับสังคม (ภาษาและอำนาจ, ภาษาและเพศสภาพ)": f"Create {num_questions} Thai sociolinguistics exercises for ม.3. Cover: language and power, language and gender, social aspects of language.",
            "วรรณคดีไทย (นิทานรามเกียรติ์, ขุนช้างขุนแผน)": f"Create {num_questions} Thai epic literature exercises for ม.3. Cover: Ramakien, Khun Chang Khun Paen, epic analysis.",
            "การอ่านวิพากษ์ (วิพากษ์บทความ, วิพากษ์ข่าว)": f"Create {num_questions} Thai critical reading exercises for ม.3. Cover: article critique, news critique, analytical reading.",
            "การเขียนวิชาการ (รายงานวิจัย, บทความวิชาการ)": f"Create {num_questions} Thai academic writing exercises for ม.3. Cover: research reports, academic articles, formal writing.",
            "วาทีวิทยา (การโต้แย้ง, การเขียนข้อเสนอ)": f"Create {num_questions} Thai rhetoric exercises for ม.3. Cover: argumentation, proposal writing, persuasive techniques.",
            # ม.4
            "ภาษากับการสื่อสาร (ภาษาในองค์กร, ภาษาธุรกิจ)": f"Create {num_questions} Thai business communication exercises for ม.4. Cover: organizational language, business Thai, professional communication.",
            "หลักภาษาไทย (ทฤษฎีภาษา, ภาษากับความคิด)": f"Create {num_questions} Thai linguistics exercises for ม.4. Cover: language theory, language and thought, linguistic concepts.",
            "วรรณคดีร่วมสมัย (นิยายไทยร่วมสมัย)": f"Create {num_questions} Thai contemporary literature exercises for ม.4. Cover: modern Thai novels, contemporary Thai fiction, literary trends.",
            "การเขียนเชิงวิชาการ (บทความวิเคราะห์)": f"Create {num_questions} Thai analytical writing exercises for ม.4. Cover: analytical essays, critical analysis, academic writing style.",
            "สื่อและภาษา (ภาษาโฆษณา, ภาษาสื่อ)": f"Create {num_questions} Thai media language exercises for ม.4. Cover: advertising language, media Thai, language in media.",
            # ม.5
            "วรรณคดีไทยและอาเซียน (วรรณคดีอาเซียน)": f"Create {num_questions} Thai and ASEAN literature exercises for ม.5. Cover: Thai literature, ASEAN literary works, regional literature.",
            "ภาษาและวัฒนธรรม (ภาษากับวัฒนธรรมไทย)": f"Create {num_questions} Thai language and culture exercises for ม.5. Cover: language and Thai culture, cultural expressions, cultural context.",
            "การนำเสนอ (การพูดในที่สาธารณะ)": f"Create {num_questions} Thai public speaking exercises for ม.5. Cover: public presentations, speech delivery, oral communication.",
            "การเขียนสร้างสรรค์ (บทละคร, บทภาพยนตร์)": f"Create {num_questions} Thai creative writing - drama and film exercises for ม.5. Cover: play scripts, screenplay writing, dramatic structure.",
            "วาทีวิทยา (การโต้วาที, การเขียนข้อเสนอ)": f"Create {num_questions} Thai debate and proposal writing exercises for ม.5. Cover: debate skills, proposal writing, persuasive discourse.",
            # ม.6
            "ภาษากับเทคโนโลยี (ภาษาอินเทอร์เน็ต, ภาษาโซเชียล)": f"Create {num_questions} Thai digital language exercises for ม.6. Cover: internet language, social media Thai, digital communication.",
            "ภาษาและอาชีพ (ภาษาสำหรับอาชีพต่างๆ)": f"Create {num_questions} Thai professional language exercises for ม.6. Cover: language for various careers, professional Thai, career communication.",
            "วรรณคดีและภาพยนตร์ (การดัดแปลงวรรณคดี)": f"Create {num_questions} Thai literature and film adaptation exercises for ม.6. Cover: novel to film adaptations, literary adaptation analysis.",
            "การเขียนเพื่อสื่อสาร (บทความสารคดี)": f"Create {num_questions} Thai documentary writing exercises for ม.6. Cover: documentary articles, feature writing, communicative writing.",
            "การประเมินผลงานภาษา (การวิจารณ์, การประเมิน)": f"Create {num_questions} Thai language evaluation exercises for ม.6. Cover: literary criticism, language assessment, evaluative writing.",
        }
        
        # Get prompt for topic or use exercise type prompt
        if topic in topic_prompts:
            base_prompt = topic_prompts[topic]
        else:
            base_prompt = exercise_prompts.get(exercise_type, exercise_prompts["mix"])
        
        full_prompt = f"""
        {base_prompt}
        
        Target Grade/Level: {grade_context.get(grade_level, grade_level)}
        Curriculum: หลักสูตรกระทรวงศึกษาธิการ (Thailand Ministry of Education)
        
        Output format:
        Q: [Question in Thai]
        A: [Answer in Thai]
        
        Create exactly {num_questions} questions.
        Include a mix of: multiple choice, fill in the blank, short answer, and true/false questions.
        Focus on: การเขียน (writing), การอ่าน (reading), หลักภาษา (grammar), คำศัพท์ (vocabulary), and วรรณคดี (literature) as appropriate.
        """
        
        try:
            resp_text = self._generate_content(full_prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    def generate_english_worksheet(self, topic, grade_level, num_questions=10, exercise_type="mix"):
        """Generates English language worksheets using AI based on English curriculum."""
        if not self.model:
            return ["Error: No AI Key"], ["Please add API Key"]
        
        # Grade context mapping
        grade_context = {
            "ป.1": "Primary 1 / Grade 1 (Elementary)",
            "ป.2": "Primary 2 / Grade 2 (Elementary)",
            "ป.3": "Primary 3 / Grade 3 (Elementary)",
            "ป.4": "Primary 4 / Grade 4 (Elementary)",
            "ป.5": "Primary 5 / Grade 5 (Elementary)",
            "ป.6": "Primary 6 / Grade 6 (Elementary)",
            "ม.1": "Matthayom 1 / Grade 7 (Lower Secondary)",
            "ม.2": "Matthayom 2 / Grade 8 (Lower Secondary)",
            "ม.3": "Matthayom 3 / Grade 9 (Lower Secondary)",
            "ม.4": "Matthayom 4 / Grade 10 (Upper Secondary)",
            "ม.5": "Matthayom 5 / Grade 11 (Upper Secondary)",
            "ม.6": "Matthayom 6 / Grade 12 (Upper Secondary)",
        }
        
        # Exercise type prompts
        exercise_prompts = {
            "mix": f"Create {num_questions} English language exercises covering various aspects including grammar, vocabulary, reading, writing, listening scripts, and speaking prompts appropriate for {grade_context.get(grade_level, grade_level)}.",
            "grammar": f"Create {num_questions} English grammar exercises for {grade_context.get(grade_level, grade_level)}. Include: fill in the blank, multiple choice, sentence transformation, and error identification.",
            "vocabulary": f"Create {num_questions} English vocabulary exercises for {grade_context.get(grade_level, grade_level)}. Include: word matching, fill in the blank, word puzzles, synonyms, antonyms, and word usage.",
            "reading": f"Create {num_questions} English reading comprehension exercises for {grade_context.get(grade_level, grade_level)}. Include: short passages, reading for understanding, and answering questions.",
            "writing": f"Create {num_questions} English writing exercises for {grade_context.get(grade_level, grade_level)}. Include: sentence writing, paragraph writing, letter writing, essay writing, and creative writing.",
            "listening": f"Create {num_questions} English listening scripts for {grade_context.get(grade_level, grade_level)}. Include: short dialogues, passages for teachers to read aloud, and comprehension questions.",
            "speaking": f"Create {num_questions} English speaking prompts/flashcards for {grade_context.get(grade_level, grade_level)}. Include: conversation starters, role-play scenarios, picture description prompts, and oral exercises.",
        }
        
        # Topic-specific prompts for better accuracy
        topic_prompts = {
            # ป.1
            "Alphabet (A-Z uppercase/lowercase)": f"Create {num_questions} English alphabet exercises for Grade 1. Cover: uppercase and lowercase letters A-Z, letter recognition, letter tracing, and alphabet order.",
            "Phonics (Aa-Zz sounds)": f"Create {num_questions} English phonics exercises for Grade 1. Cover: letter sounds, initial sounds, phonetic awareness, and sound-letter correspondence.",
            "Numbers 1-10 (counting)": f"Create {num_questions} English numbers 1-10 exercises for Grade 1. Cover: number words, counting, number recognition, and simple number sentences.",
            "Colors (Red, blue, green, yellow, etc.)": f"Create {num_questions} English colors vocabulary exercises for Grade 1. Cover: color names, color identification, color matching, and using colors in sentences.",
            "Shapes (Circle, square, triangle, etc.)": f"Create {num_questions} English shapes vocabulary exercises for Grade 1. Cover: shape names (circle, square, triangle, rectangle, star, heart), shape identification, and properties.",
            "Body Parts (Head, eyes, ears, nose, etc.)": f"Create {num_questions} English body parts vocabulary exercises for Grade 1. Cover: body part names, body part identification, and simple sentences about body parts.",
            "Family (Mother, father, sister, brother)": f"Create {num_questions} English family vocabulary exercises for Grade 1. Cover: family members, family relationships, and introducing family.",
            "Animals (Cat, dog, bird, fish, etc.)": f"Create {num_questions} English animal vocabulary exercises for Grade 1. Cover: animal names, animal sounds, animal identification, and animal facts.",
            
            # ป.2
            "Numbers 11-100 (counting)": f"Create {num_questions} English numbers 11-100 exercises for Grade 2. Cover: number words, counting by tens, place value, and number patterns.",
            "Days & Months (Monday-Sunday, Jan-Dec)": f"Create {num_questions} English days and months exercises for Grade 2. Cover: days of the week, months of the year, calendar vocabulary, and date writing.",
            "Time (O'clock, half past)": f"Create {num_questions} English time exercises for Grade 2. Cover: telling time (o'clock, half past), clock reading, and time vocabulary.",
            "Food & Drinks (Rice, bread, water, milk)": f"Create {num_questions} English food and drinks vocabulary exercises for Grade 2. Cover: food names, drink names, healthy eating, and food categories.",
            "Clothing (Shirt, pants, dress, shoes)": f"Create {num_questions} English clothing vocabulary exercises for Grade 2. Cover: clothing items, clothing descriptions, weather and clothing.",
            "Weather (Hot, cold, rainy, sunny)": f"Create {num_questions} English weather vocabulary exercises for Grade 2. Cover: weather words, weather descriptions, and weather forecasting.",
            "Places (School, home, market, hospital)": f"Create {num_questions} English places vocabulary exercises for Grade 2. Cover: common places, directions to places, and community buildings.",
            "Greetings (Hello, goodbye, thank you)": f"Create {num_questions} English greetings and polite expressions exercises for Grade 2. Cover: greetings, farewells, polite words, and social expressions.",
            
            # ป.3
            "Present Simple (I am, you are, he/she is)": f"Create {num_questions} English Present Simple tense exercises for Grade 3. Cover: am/is/are usage, affirmative, negative, and interrogative forms.",
            "This-That-These-Those": f"Create {num_questions} English demonstratives exercises for Grade 3. Cover: this/that/these/those, singular and plural, and near/far distinction.",
            "Have-Has (possession)": f"Create {num_questions} English have/has exercises for Grade 3. Cover: possession with have/has, affirmative, negative, and questions.",
            "Prepositions (In, on, under, behind)": f"Create {num_questions} English preposition exercises for Grade 3. Cover: position prepositions (in, on, under, behind, in front of, next to) and location.",
            "WH-Questions (What, Where, When, Why, Who)": f"Create {num_questions} English WH-questions exercises for Grade 3. Cover: what, where, when, why, who questions and answering techniques.",
            "Daily Routines (Wake up, eat breakfast)": f"Create {num_questions} English daily routine exercises for Grade 3. Cover: daily activities, time expressions, and describing routines.",
            "Occupations (Doctor, teacher, farmer)": f"Create {num_questions} English occupations vocabulary exercises for Grade 3. Cover: job names, job descriptions, and what people do.",
            "Adjectives (Big, small, tall, beautiful)": f"Create {num_questions} English adjectives exercises for Grade 3. Cover: descriptive adjectives, comparison, and using adjectives in sentences.",
            
            # ป.4
            "Past Simple (was/were)": f"Create {num_questions} English Past Simple (was/were) exercises for Grade 4. Cover: affirmative, negative, questions with was/were, and time expressions.",
            "Regular Verbs (Played, watched, cleaned)": f"Create {num_questions} English regular past tense verbs exercises for Grade 4. Cover: -ed endings, spelling rules, and affirmative sentences.",
            "Irregular Verbs (Went, ate, drank, saw)": f"Create {num_questions} English irregular past tense verbs exercises for Grade 4. Cover: common irregular verbs, past forms, and sentence building.",
            "Object Pronouns (Me, him, her, us, them)": f"Create {num_questions} English object pronouns exercises for Grade 4. Cover: object pronouns, subject vs object, and usage in sentences.",
            "There is-There are": f"Create {num_questions} English there is/there are exercises for Grade 4. Cover: affirmative, negative, questions, and singular/plural.",
            "Commands (Open the door, close the window)": f"Create {num_questions} English imperative/commands exercises for Grade 4. Cover: giving commands, instructions, and imperatives.",
            "Descriptions (Describing people/things)": f"Create {num_questions} English description exercises for Grade 4. Cover: describing appearance, personality, and objects.",
            "School Subjects (Math, English, Science, Art)": f"Create {num_questions} English school subjects vocabulary exercises for Grade 4. Cover: subject names, schedules, and preferences.",
            
            # ป.5
            "Future Will-Going to": f"Create {num_questions} English future tense (will vs going to) exercises for Grade 5. Cover: predictions, plans, intentions, and differences between will and going to.",
            "Present Continuous (am/is/are + verb-ing)": f"Create {num_questions} English Present Continuous tense exercises for Grade 5. Cover: affirmative, negative, questions, and present vs past usage.",
            "Can-Could (ability, permission)": f"Create {num_questions} English can/could exercises for Grade 5. Cover: ability, permission, requests, and polite expressions.",
            "Some-Any": f"Create {num_questions} English some/any exercises for Grade 5. Cover: countable and uncountable nouns, affirmative, negative, and questions.",
            "Telling Time (Quarter past, quarter to)": f"Create {num_questions} English advanced time telling exercises for Grade 5. Cover: quarter past, quarter to, past/to expressions, and schedules.",
            "Giving Directions (Turn left, turn right)": f"Create {num_questions} English giving directions exercises for Grade 5. Cover: direction words, giving directions, and following maps.",
            "Invitations (Would you like...?, Let's...)": f"Create {num_questions} English invitation exercises for Grade 5. Cover: making invitations, accepting, declining, and suggestions.",
            "Letter Writing (Formal and informal)": f"Create {num_questions} English letter writing exercises for Grade 5. Cover: formal letters, informal letters, parts of a letter, and writing practice.",
            
            # ป.6
            "Tenses Review (Present, Past, Future)": f"Create {num_questions} English tenses review exercises for Grade 6. Cover: all three tenses, usage, and sentence transformation.",
            "Modal Verbs (Must, should, have to, may)": f"Create {num_questions} English modal verbs exercises for Grade 6. Cover: must, should, have to, may, might, can, could - meanings and usage.",
            "Passive Voice (is/are + verb3)": f"Create {num_questions} English passive voice exercises for Grade 6. Cover: present/past passive, affirmative, negative, and questions.",
            "If Clauses (Conditionals type 1)": f"Create {num_questions} English conditionals type 1 exercises for Grade 6. Cover: if clauses, real conditions, and cause-effect sentences.",
            "Reported Speech (Said, told, asked)": f"Create {num_questions} English reported speech exercises for Grade 6. Cover: direct vs reported speech, tense changes, and pronoun changes.",
            "Conjunctions (And, but, or, because, so)": f"Create {num_questions} English conjunctions exercises for Grade 6. Cover: coordinating conjunctions, compound sentences, and usage.",
            "Reading Comprehension (Passages, questions)": f"Create {num_questions} English reading comprehension exercises for Grade 6. Cover: various passages, comprehension questions, and reading strategies.",
            "Paragraph Writing (3-5 sentences)": f"Create {num_questions} English paragraph writing exercises for Grade 6. Cover: topic sentence, supporting details, concluding sentence.",
            
            # ม.1
            "Present Perfect (have/has + verb3)": f"Create {num_questions} English Present Perfect exercises for ม.1. Cover: affirmative, negative, questions, ever, never, for, since.",
            "Since-For (time expressions)": f"Create {num_questions} English since/for time expressions exercises for ม.1. Cover: duration vs point in time, present perfect usage.",
            "Tag Questions (aren't you?, isn't it?)": f"Create {num_questions} English tag questions exercises for ม.1. Cover: forming tag questions, affirmative/negative tags, and responses.",
            "Relative Clauses (Who, which, that)": f"Create {num_questions} English relative clauses exercises for ม.1. Cover: defining relative clauses, who, which, that, and relative pronouns.",
            "Gerunds & Infinitives": f"Create {num_questions} English gerunds and infinitives exercises for ม.1. Cover: verb + ing, verb + to, and usage patterns.",
            "Making Suggestions (Let's, Why don't we)": f"Create {num_questions} English making suggestions exercises for ม.1. Cover: Let's, Why don't we, How about, I suggest.",
            "Phone Conversations": f"Create {num_questions} English phone conversation exercises for ม.1. Cover: answering phone, leaving messages, phone etiquette.",
            "Shopping & Money": f"Create {num_questions} English shopping and money exercises for ม.1. Cover: prices, bargaining, making purchases, and transactions.",
            "Travel & Transportation": f"Create {num_questions} English travel and transportation exercises for ม.1. Cover: modes of transport, booking tickets, travel vocabulary.",
            "Health & Fitness": f"Create {num_questions} English health and fitness exercises for ม.1. Cover: health problems, doctor visits, healthy lifestyle.",
            
            # ม.2
            "Past Continuous (was/were + verb-ing)": f"Create {num_questions} English Past Continuous exercises for ม.2. Cover: affirmative, negative, questions, and Past Simple vs Continuous.",
            "Future Continuous (will be + verb-ing)": f"Create {num_questions} English Future Continuous exercises for ม.2. Cover: will be + verb-ing, time expressions, and usage.",
            "Conditionals Type 2 (If I were, I would)": f"Create {num_questions} English conditionals type 2 exercises for ม.2. Cover: if + past simple, would + verb, hypothetical situations.",
            "Reported Questions": f"Create {num_questions} English reported questions exercises for ม.2. Cover: direct vs reported questions, question words in reported speech.",
            "Quantifiers (Much, many, a few, a little)": f"Create {num_questions} English quantifiers exercises for ม.2. Cover: much, many, a few, a little, some, any, and usage with nouns.",
            "Comparison (Adjectives, adverbs)": f"Create {num_questions} English comparison exercises for ม.2. Cover: comparative and superlative forms, comparison structures.",
            "Wish Sentences (I wish I could...)": f"Create {num_questions} English wish sentences exercises for ม.2. Cover: wishes about present, wishes about past, and regrets.",
            "Email Writing (Formal and informal)": f"Create {num_questions} English email writing exercises for ม.2. Cover: formal emails, informal emails, email structure.",
            "News Writing": f"Create {num_questions} English news writing exercises for ม.2. Cover: news article structure, headline writing, 5 Ws.",
            "Story Writing": f"Create {num_questions} English story writing exercises for ม.2. Cover: narrative structure, story elements, creative writing.",
            
            # ม.3
            "Conditionals All Types (Type 1, 2, 3)": f"Create {num_questions} English all types of conditionals exercises for ม.3. Cover: Type 1, 2, 3, mixed conditionals.",
            "Passive Voice (All tenses)": f"Create {num_questions} English all tenses passive voice exercises for ม.3. Cover: all tenses in passive, by + agent, and sentence transformation.",
            "Reported Speech (All reporting verbs)": f"Create {num_questions} English all reported speech exercises for ม.3. Cover: say, tell, ask, and other reporting verbs, tense changes.",
            "Gerunds & Infinitives (Special uses)": f"Create {num_questions} English advanced gerunds and infinitives exercises for ม.3. Cover: special uses, patterns, and exceptions.",
            "Modal Perfects (Should have, could have)": f"Create {num_questions} English modal perfect exercises for ม.3. Cover: should have, could have, would have, might have.",
            "Articles (A, an, the, zero article)": f"Create {num_questions} English articles exercises for ม.3. Cover: a, an, the, zero article, and specific rules.",
            "Essay Writing (Opinion, comparison)": f"Create {num_questions} English essay writing exercises for ม.3. Cover: opinion essays, comparison essays, essay structure.",
            "O-NET Preparation (Grammar, vocabulary)": f"Create {num_questions} English O-NET preparation exercises for ม.3. Cover: grammar review, vocabulary, test strategies.",
            "Critical Reading": f"Create {num_questions} English critical reading exercises for ม.3. Cover: analyzing texts, inference, evaluation, and critical thinking.",
            "Creative Writing": f"Create {num_questions} English creative writing exercises for ม.3. Cover: creative expression, storytelling techniques, descriptive writing.",
            
            # ม.4
            "Narrative Tenses (Past perfect)": f"Create {num_questions} English narrative tenses exercises for ม.4. Cover: past perfect, past perfect continuous, story sequencing.",
            "Future Perfect (will have + verb3)": f"Create {num_questions} English Future Perfect exercises for ม.4. Cover: will have + past participle, by + time expressions.",
            "Mixed Conditionals": f"Create {num_questions} English mixed conditionals exercises for ม.4. Cover: combining different conditional types, advanced structures.",
            "Wish-Remorse (I wish I had...)": f"Create {num_questions} English wish and remorse exercises for ม.4. Cover: wishes about past, regrets, I wish + past perfect.",
            "Linking Words (However, although, despite)": f"Create {num_questions} English linking words exercises for ม.4. Cover: however, although, despite, moreover, therefore, and connectors.",
            "Paragraph Development": f"Create {num_questions} English paragraph development exercises for ม.4. Cover: topic sentence, supporting details, coherence, unity.",
            "Speaking: Opinions (I think, In my opinion)": f"Create {num_questions} English expressing opinions exercises for ม.4. Cover: I think, in my opinion, I believe, discussing views.",
            "Vocabulary 1500 (Word families, synonyms)": f"Create {num_questions} English vocabulary building exercises for ม.4. Cover: word families, synonyms, antonyms, collocations.",
            "Academic Vocabulary": f"Create {num_questions} English academic vocabulary exercises for ม.4. Cover: academic words, formal language, scholarly expressions.",
            "Debating Skills": f"Create {num_questions} English debating skills exercises for ม.4. Cover: arguments, counterarguments, debate structure, presentation.",
            
            # ม.5
            "Mixed Tenses Review": f"Create {num_questions} English mixed tenses review exercises for ม.5. Cover: all tenses in context, tense selection, and accuracy.",
            "Modal Verbs Review (Must, have to, should)": f"Create {num_questions} English modal verbs review exercises for ม.5. Cover: must, have to, should, ought to, mustn't, don't have to.",
            "Participle Clauses": f"Create {num_questions} English participle clauses exercises for ม.5. Cover: present participles, past participles, reduced clauses.",
            "Passive Voice Review": f"Create {num_questions} English passive voice review exercises for ม.5. Cover: all tenses, formal writing, emphasis, and transformation.",
            "Essay Types (Argumentative, descriptive)": f"Create {num_questions} English essay types exercises for ม.5. Cover: argumentative essays, descriptive essays, structure and techniques.",
            "Speaking: Debating (Agree/disagree)": f"Create {num_questions} English debating and disagreement exercises for ม.5. Cover: agreeing, disagreeing, giving reasons, persuasion.",
            "Listening Skills (News, interviews)": f"Create {num_questions} English listening skills exercises for ม.5. Cover: news listening, interview comprehension, note-taking.",
            "Vocabulary 2000 (Idioms, phrasal verbs)": f"Create {num_questions} English advanced vocabulary exercises for ม.5. Cover: idioms, phrasal verbs, expressions, and usage.",
            "Academic Writing": f"Create {num_questions} English academic writing exercises for ม.5. Cover: research writing, citations, formal style, academic conventions.",
            "Presentation Skills": f"Create {num_questions} English presentation skills exercises for ม.5. Cover: presentation structure, delivery, visual aids, handling questions.",
            
            # ม.6
            "Advanced Grammar (Inversion, emphasis)": f"Create {num_questions} English advanced grammar exercises for ม.6. Cover: inversion, emphasis, cleft sentences, advanced structures.",
            "Academic Writing (Research, citations)": f"Create {num_questions} English academic writing exercises for ม.6. Cover: research papers, citations, bibliography, academic integrity.",
            "Critical Reading (Analysis, inference)": f"Create {num_questions} English critical reading exercises for ม.6. Cover: text analysis, inference, evaluation, literary criticism.",
            "Presentation Skills": f"Create {num_questions} English formal presentation exercises for ม.6. Cover: presentations, public speaking, professional communication.",
            "Test Preparation (O-NET, University entrance)": f"Create {num_questions} English test preparation exercises for ม.6. Cover: O-NET, university entrance exam practice, strategies.",
            "Career English (Resume, interview)": f"Create {num_questions} English career English exercises for ม.6. Cover: resume writing, job interviews, cover letters, professional communication.",
            "Global Issues (Environment, technology)": f"Create {num_questions} English global issues exercises for ม.6. Cover: environment, technology, climate change, social issues.",
            "Literature (Poems, short stories)": f"Create {num_questions} English literature exercises for ม.6. Cover: poetry analysis, short stories, literary devices, appreciation.",
        }
        
        # Get prompt for topic or use exercise type prompt
        if topic in topic_prompts:
            base_prompt = topic_prompts[topic]
        else:
            base_prompt = exercise_prompts.get(exercise_type, exercise_prompts["mix"])
        
        full_prompt = f"""
        {base_prompt}
        
        Target Grade/Level: {grade_context.get(grade_level, grade_level)}
        Curriculum: Thailand English Language Curriculum (กระทรวงศึกษาธิการ)
        
        Output format:
        Q: [Question in English]
        A: [Answer in English]
        
        Create exactly {num_questions} questions.
        Include a mix of: multiple choice, fill in the blank, short answer, matching, and true/false questions as appropriate.
        Focus on: Grammar, Vocabulary, Reading, Writing, Listening scripts, and Speaking prompts as selected.
        Make exercises appropriate for the grade level and engaging for students.
        """
        
        try:
            resp_text = self._generate_content(full_prompt)
            qs, ans = [], []
            for line in resp_text.split('\n'):
                if line.startswith("Q:"): qs.append(line.replace("Q:", "").strip())
                elif line.startswith("A:"): ans.append(line.replace("A:", "").strip())
            return qs, ans
        except Exception as e:
            return [f"AI Error: {str(e)}"], ["Error"]

    # --- PDF Creation ---
    def create_qr_code(self, data):
        qr = qrcode.QRCode(box_size=10, border=1)
        qr.add_data(data); qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        b = io.BytesIO(); img.save(b, format="PNG"); b.seek(0)
        return b

    def create_pdf(self, title, school_name, content_type, data, answers=None, qr_link=None, uploaded_logo=None):
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        w, h = A4
        
        # Header
        c.setStrokeColor(blue); c.setLineWidth(2); c.setDash(6,3)
        c.rect(1*cm, 1*cm, w-2*cm, h-2*cm); c.setDash([]); c.setStrokeColor(black)
        
        if uploaded_logo:
            try:
                logo = Image.open(uploaded_logo)
                c.drawImage(ImageReader(logo), 2*cm, h-4*cm, width=2.5*cm, height=2.5*cm, mask='auto')
            except: pass

        c.setFont(self.font_name, 24); c.drawCentredString(w/2, h-2.5*cm, title)
        c.setFont(self.font_name, 14); c.setFillColor(blue); c.drawCentredString(w/2, h-3.5*cm, school_name); c.setFillColor(black)
        c.setFont(self.font_name, 12)
        c.drawString(2*cm, h-4.5*cm, "Name: __________________________  Date: ______________")

        if qr_link:
            qr = self.create_qr_code(qr_link)
            c.drawImage(ImageReader(qr), w-4*cm, h-3.5*cm, width=2.5*cm, height=2.5*cm)

        # Body
        y = h - 6*cm
        
        # Handle all content types (including IPST topics)
        if content_type in ["Word Search"]:
            grid, words = data
            c.setFont("Courier-Bold", 14)
            start_x = 3*cm
            cell = 0.8*cm
            for r, row in enumerate(grid):
                for col, char in enumerate(row):
                    c.drawString(start_x + col*cell, y - r*cell, char)
            
            y_words = y - len(grid)*cell - 1*cm
            c.setFont(self.font_name, 12)
            c.drawString(2*cm, y_words, "Find: " + ", ".join(words))
        
        elif content_type in ["Handwriting Practice"]:
            c.setFont("Courier", 30); c.setFillColor(lightgrey)
            for line in data:
                c.setLineWidth(1); c.setStrokeColor(black)
                c.line(2*cm, y+10, w-2*cm, y+10)
                c.line(2*cm, y-10, w-2*cm, y-10)
                c.setDash(3,3); c.line(2*cm, y, w-2*cm, y); c.setDash([])
                c.drawString(2*cm, y-8, line)
                y -= 2.5*cm
                if y < 3*cm: c.showPage(); y = h-5*cm
        
        else:
            # Default: render as questions/worksheet (all IPST topics)
            c.setFont(self.font_name, 14)
            for i, q in enumerate(data):
                # Basic wrapping logic
                text = c.beginText(2.5*cm, y)
                text.setFont(self.font_name, 14)
                # Split newlines for multi-line questions
                lines = q.split('\n')
                for line in lines:
                    text.textLine(line)
                c.drawText(text)
                
                # Dynamic spacing
                height_needed = len(lines) * 0.7*cm + 1*cm
                y -= height_needed
                if y < 3*cm: c.showPage(); y = h-4*cm

        c.showPage()
        
        # Answers
        if answers:
            c.setFont(self.font_name, 20); c.drawCentredString(w/2, h-2.5*cm, f"ANSWER KEY: {title}")
            y = h-4*cm; c.setFont(self.font_name, 12)
            for i, ans in enumerate(answers):
                c.drawString(3*cm, y, f"{i+1}) {ans}"); y -= 0.8*cm
                if y < 2*cm: c.showPage(); y = h-4*cm
            c.showPage()

        c.save(); buffer.seek(0)
        return buffer

    # --- Word Doc ---
    def create_word_doc(self, title, school_name, content_type, data, answers=None):
        doc = Document()
        # Header/Setup
        doc.add_heading(title, 0)
        doc.add_paragraph(f"{school_name}\nName: ________________ Date: _________")
        
        # Handle all content types
        if content_type in ["Word Search"]:
            grid, words = data
            table = doc.add_table(len(grid), len(grid[0]))
            for r, row in enumerate(grid):
                for c, char in enumerate(row):
                    table.cell(r,c).text = char
            doc.add_paragraph("Find: " + ", ".join(words))
        elif content_type in ["Handwriting Practice"]:
            for line in data: doc.add_paragraph(line)
        else:
            # Default: render as questions/worksheet (all IPST topics)
            for i, q in enumerate(data):
                doc.add_paragraph(f"{i+1}. {q}")
                doc.add_paragraph("_"*20)
            
        if answers:
            doc.add_page_break()
            doc.add_heading("Answers", 1)
            for i, a in enumerate(answers): doc.add_paragraph(f"{i+1}) {a}")
            
        b = io.BytesIO(); doc.save(b); b.seek(0)
        return b
